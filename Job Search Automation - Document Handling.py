import os
import win32com.client
import re
from distutils.dir_util import copy_tree
from pathlib import Path
from pywebcopy import save_webpage
from docx import Document

#Enter Company name, Title and the url of the jobposting
Company = 'DR'
Title = 'Journalist'
url = 'https://candidate.hr-manager.net/ApplicationInit.aspx?cid=1165&ProjectId=146575&DepartmentId=18960&MediaId=59'
Newdir = r"C:\Users\chris\OneDrive\Python\Jobsøgnings automatisering"+ '\\' + Company + '\\' + Title

#Create Directory with Company name if it does not exist
Path(Newdir ).mkdir(parents=True, exist_ok=True)
print('Directory Created')

#Copy Standard folder
Directory = r'C:\Users\chris\OneDrive\Python\Jobsøgnings automatisering'
fromDirectory = r'C:\Users\chris\OneDrive\Python\Jobsøgnings automatisering\Standard folder'
toDirectory = Directory + '\\' + Company +'\\'+ Title
copy_tree(fromDirectory, toDirectory)
print('Contents copied')

#Open document for editing
document = Document(r'C:\Users\chris\OneDrive\Python\Jobsøgnings automatisering\DR\Journalist\Christoffer Mønster Frydkjær_Cover letter.docx')

#Find and replace word
for run in [run for par in document.paragraphs for run in par.runs] + \
           [run for table in document.tables for col in table.columns for cell in col.cells for par in cell.paragraphs for run in par.runs]:
    s = run.text.replace("____", "DR")
    if s != run.text:    # rewrite run.text if (and only if) it has changed; *always* rewriting is not good, it could destroy column breaks, etc.
        run.text = s

#Save document        
document.save(r'C:\Users\chris\OneDrive\Python\Jobsøgnings automatisering\DR\Journalist\Christoffer Mønster Frydkjær_Cover letter.docx')
print('Document edited and saved')

#Convert word to file to pdf
path = (r'C:\Users\chris\OneDrive\Python\Jobsøgnings automatisering'+ '\\'  + Company + '\\' + Title)
word_file_names = []
word = win32com.client.Dispatch('Word.Application')
for dirpath, dirnames, filenames in os.walk(path):
    for f in filenames:  
        if f.lower().endswith(".docx") :
            new_name = f.replace(".docx", ".pdf")
            in_file =(dirpath + '/'+ f)
            new_file =(dirpath + '/' + new_name)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(new_file, FileFormat = 17)
            doc.Close()
        if f.lower().endswith(".doc"):
            new_name = f.replace(".doc", ".pdf")
            in_file =(dirpath +'/' + f)
            new_file =(dirpath +'/' + new_name)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(new_file, FileFormat = 17)
            doc.Close()
word.Quit()
print('Word files converted to pdf')

#Create Jobposting file for record saved html and java
kwargs = {'bypass_robots': True, 'project_name': 'Jobposting'}
save_webpage(url, Newdir, **kwargs)
print('Jobposting saved')








